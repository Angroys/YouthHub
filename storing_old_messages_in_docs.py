import pytz, os, sys, logging, time, asyncio

sys.path.append(os.path.abspath(os.path.dirname('python_dependencies\groq_api.py')))

from datetime import datetime
from docx import Document
from python_dependencies.telegram_api_code import TelegramClass
from python_dependencies.groq_api import GroqAPI


utc=pytz.UTC

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration
grok_api_key = ''
api_id = 
api_hash = ''
download_path = 'download/'
telegram_channel = ''  # The name of the Telegram channel
limit = 1000  # Number of past messages to fetch (increase this if necessary)

# Define the date range (start and end date for fetching messages)
start_date = utc.localize(datetime(2024, 1, 1))  # Start date (YYYY, MM, DD)
end_date = utc.localize(datetime(2024, 12, 30))  # End date (YYYY, MM, DD)

# Initialize Telegram and Groq clients
telegram_client = TelegramClass(api_id, api_hash, download_path, groq_api_key=grok_api_key)
groq_client = GroqAPI(grok_api_key)

# Function to save messages to a DOCX file
def save_messages_to_docx(messages, file_path):
    try:
        # Create a new Document
        doc = Document()
        doc.add_heading(f'Messages from {start_date.strftime("%Y/%m/%d")} to {end_date.strftime("%Y-%m-%d")}', 0)


        for message in messages:
            doc.add_paragraph(str(groq_client.customMessage("Extract text in original form, and the link too. without additional commentaries from your part", message)))
            logger.info(f"Message added")
        # Save the document
        doc.save(file_path)
        logger.info(f"Messages saved to {file_path}")
    except Exception as e:
        logger.error(f"Error saving messages to DOCX: {e}")

# Function to fetch messages from a specific date range from a Telegram channel
async def fetch_messages_in_date_range(chat_id, start_date, end_date, limit=1000):
    messages = []
    try:
        # Fetch past messages and filter based on the date range
        async for message in telegram_client.client.iter_messages(chat_id, limit=limit):
            if start_date <= message.date <= end_date:
                messages.append(message)
                logger.info(f"Message ID: {message.id}, Sender ID: {message.sender_id}, Date: {message.date}, Text: {message.text}")
            elif message.date < start_date:
                # Stop fetching if the message date is before the start date
                break
    except Exception as e:
        logger.error(f"Error fetching messages from {chat_id}: {e}")
    return messages

# Main function to run the bot
async def run_bot():
    try:
        await telegram_client.connect()  # Ensure connection
        logger.info(f"Bot is connected and fetching messages from {telegram_channel} between {start_date.strftime('%Y-%m-%d')} and {end_date.strftime('%Y-%m-%d')}...")

        # Fetch messages from the Telegram channel within the date range
        messages = await fetch_messages_in_date_range(telegram_channel, start_date, end_date, limit)

        # Filter messages using GroqAPI
        filtered_messages = []
        for message in messages:
            try:
                logger.info("Searching for Messages")
                if groq_client.messageIsAListoFnEWS(message.text) == 'True':
                    filtered_messages.append(message)
            except Exception as e:
                logger.error(f"Error processing message in chat {telegram_channel}: {e}")


        # Save the filtered messages to a DOCX file
        if filtered_messages:
            docx_file_path = os.path.join(download_path, f"messages_{start_date.strftime('%Y.%m.%d')}_to_{end_date.strftime('%Y%m%d')}.docx")
            save_messages_to_docx(filtered_messages, docx_file_path)
        else:
            logger.info("No relevant messages found.")
    except Exception as e:
        logger.error(f"An error occurred: {e}")
    finally:
        await telegram_client.disconnect()  # Disconnect after all tasks are complete

if __name__ == "__main__":
    try:
        # Ensure the download directory exists
        if not os.path.exists(download_path):
            os.makedirs(download_path)

        asyncio.run(run_bot())  # Run the bot asynchronously
    except Exception as e:
        logger.error(f"An error occurred in the main execution: {e}")
